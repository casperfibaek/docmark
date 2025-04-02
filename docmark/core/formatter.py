"""
Formatter module for DocMark.

This module provides functionality for formatting Markdown and DOCX documents,
ensuring consistent and pretty output.
"""

from typing import Optional

from docmark.utils import text as text_utils
from docmark.core.llm import LLMManager


class MarkdownFormatter:
    """Formatter for Markdown documents."""

    def __init__(self, llm_manager: Optional[LLMManager] = None):
        """
        Initialize the Markdown formatter.

        Parameters
        ----------
        llm_manager : LLMManager, optional
            LLM manager for enhanced formatting.
        """
        self.llm_manager = llm_manager

    def format(self, markdown_content: str, use_llm: bool = True) -> str:
        """
        Format Markdown content.

        Parameters
        ----------
        markdown_content : str
            The Markdown content to format.
        use_llm : bool, optional
            Whether to use LLM for enhanced formatting.

        Returns
        -------
        str
            Formatted Markdown content.
        """
        # Save the original content for verification
        original_content = markdown_content

        # Apply basic formatting fixes
        basic_formatted = text_utils.fix_markdown_formatting(markdown_content)

        # Replace TOC marker with generated TOC
        basic_formatted = text_utils.replace_toc_marker(basic_formatted)

        # Use LLM for enhanced formatting if available
        if use_llm and self.llm_manager:
            try:
                # The enhanced LLM module now handles chunking for large documents automatically
                llm_formatted = self.llm_manager.enhance_markdown(basic_formatted)

                # Verify the output is not truncated (comparing character counts allowing for some formatting changes)
                original_length = len(basic_formatted.strip())
                formatted_length = len(llm_formatted.strip())

                # Check for significant loss of content
                if formatted_length < original_length * 0.9:  # Allow for up to 10% reduction due to whitespace changes
                    print(f"Warning: LLM formatting may have truncated content. Original: {original_length} chars, Formatted: {formatted_length} chars")
                    print("Falling back to basic formatting")
                    return basic_formatted

                # Verify key content is preserved by checking for the last paragraph
                # Get the last 50 non-empty characters from the original content
                # This will help verify that the end of the document hasn't been truncated
                def get_last_paragraph(text):
                    # Split by double newlines to get paragraphs
                    paragraphs = [p for p in text.split('\n\n') if p.strip()]
                    if paragraphs:
                        # Get the last non-empty paragraph
                        return paragraphs[-1].strip()
                    return ""

                last_original_paragraph = get_last_paragraph(original_content)

                # If the last paragraph is more than 20 characters, check if a portion of it is in the formatted result
                if len(last_original_paragraph) > 20:
                    # Use a portion of the last paragraph (to account for formatting changes)
                    # Use at least 15 chars or 1/3 of the paragraph, whichever is greater
                    check_length = max(15, len(last_original_paragraph) // 3)
                    check_text = last_original_paragraph[:check_length]

                    if check_text not in llm_formatted:
                        print(f"Warning: Last paragraph not found in LLM formatted output. Falling back to basic formatting.")
                        return basic_formatted

                return llm_formatted
            except Exception as e:
                print(f"Warning: LLM formatting failed: {str(e)}")
                return basic_formatted

        return basic_formatted

    def format_table(self, html_table: str) -> str:
        """
        Format an HTML table as Markdown.

        Parameters
        ----------
        html_table : str
            The HTML table to format.

        Returns
        -------
        str
            Formatted Markdown table.
        """
        if self.llm_manager:
            try:
                return self.llm_manager.improve_table_conversion(html_table)
            except Exception as e:
                print(f"Warning: LLM table formatting failed: {str(e)}")

        # Fallback to basic table formatting
        from bs4 import BeautifulSoup, Tag


        soup = BeautifulSoup(html_table, 'html.parser')
        table = soup.find('table')

        if not isinstance(table, Tag): # Check if it's a Tag before using find_all
            return ""

        rows = table.find_all('tr')
        if not rows:
            return ""

        # Process header row
        header_row = rows[0]
        header_cells = header_row.find_all(['th', 'td'])
        header = "| " + " | ".join([cell.get_text().strip() for cell in header_cells]) + " |"

        # Create separator row
        separator = "| " + " | ".join(["---" for _ in header_cells]) + " |"

        # Process data rows
        data_rows = []
        for row in rows[1:]:
            cells = row.find_all('td')
            data_rows.append("| " + " | ".join([cell.get_text().strip() for cell in cells]) + " |")

        # Combine all rows
        return header + "\n" + separator + "\n" + "\n".join(data_rows)

    def format_code_block(self, code: str, language: str = "") -> str:
        """
        Format a code block.

        Parameters
        ----------
        code : str
            The code to format.
        language : str, optional
            The programming language of the code.

        Returns
        -------
        str
            Formatted code block.
        """
        # Ensure language is specified
        if not language:
            language = "text"

        # Format the code block
        return f"```{language}\n{code}\n```"

    def format_image(self, alt_text: str, url: str, title: Optional[str] = None) -> str:
        """
        Format an image.

        Parameters
        ----------
        alt_text : str
            The alt text for the image.
        url : str
            The URL of the image.
        title : str, optional
            The title of the image.

        Returns
        -------
        str
            Formatted image.
        """
        if title:
            return f'![{alt_text}]({url} "{title}")'
        else:
            return f'![{alt_text}]({url})'


class DocxFormatter:
    """Formatter for DOCX documents."""

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the DOCX formatter.

        Parameters
        ----------
        template_path : str, optional
            Path to the template file.
        """
        self.template_path = template_path

    def apply_heading_style(self, paragraph, level: int, doc) -> None:
        """
        Apply heading style to a paragraph.

        Parameters
        ----------
        paragraph : Paragraph
            The paragraph to style.
        level : int
            The heading level (1-6).
        doc : Document
            The document containing the paragraph.
        """
        # Map heading level to style name
        style_map = {
            1: 'Title',
            2: 'Heading 1',
            3: 'Heading 2',
            4: 'Heading 3',
            5: 'Heading 4',
            6: 'Heading 5',
        }

        # Get style name for this level
        style_name = style_map.get(level)

        # Check if style exists in document
        if style_name and style_exists(doc, style_name):
            paragraph.style = doc.styles[style_name]
        else:
            # Fallback to built-in heading styles
            paragraph.style = f'Heading {level}'

        # Set outline level for TOC
        set_outline_level(paragraph, level - 1)

    def apply_list_style(self, paragraph, is_numbered: bool, level: int, doc) -> None:
        """
        Apply list style to a paragraph.

        Parameters
        ----------
        paragraph : Paragraph
            The paragraph to style.
        is_numbered : bool
            Whether the list is numbered.
        level : int
            The list level (1-based).
        doc : Document
            The document containing the paragraph.
        """
        # Map list type and level to style name
        if is_numbered:
            style_name = f'List Number {level}' if level > 1 else 'List Number'
        else:
            style_name = f'List Bullet {level}' if level > 1 else 'List Bullet'

        # Check if style exists in document
        if style_exists(doc, style_name):
            paragraph.style = doc.styles[style_name]
        else:
            # Fallback to manual formatting
            from docx.shared import Pt

            # Add bullet or number
            if is_numbered:
                paragraph.add_run(f"{level}. ").bold = True
            else:
                paragraph.add_run("• ").bold = True

            # Add indentation for nested levels
            paragraph.paragraph_format.left_indent = Pt(36 * level)

    def apply_table_style(self, table, doc) -> None:
        """
        Apply table style to a table.

        Parameters
        ----------
        table : Table
            The table to style.
        doc : Document
            The document containing the table.
        """
        # Check if Table Grid style exists
        if style_exists(doc, 'Table Grid'):
            table.style = doc.styles['Table Grid']

        # Apply header row formatting
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    def apply_image_style(self, image, width_inches: float = 6.0) -> None:
        """
        Apply image style.

        Parameters
        ----------
        image : InlineShape
            The image to style.
        width_inches : float, optional
            The width of the image in inches.
        """
        from docx.shared import Inches

        # Set image width while maintaining aspect ratio
        image.width = Inches(width_inches)

        # Height will adjust automatically to maintain aspect ratio

    def apply_code_style(self, paragraph, doc) -> None:
        """
        Apply code style to a paragraph.

        Parameters
        ----------
        paragraph : Paragraph
            The paragraph to style.
        doc : Document
            The document containing the paragraph.
        """
        # Check if Code style exists
        if style_exists(doc, 'Code'):
            paragraph.style = doc.styles['Code']
        else:
            # Fallback to basic code formatting
            for run in paragraph.runs:
                run.font.name = 'Courier New'

            # Add light gray background
            from docx.shared import Pt
            paragraph.paragraph_format.left_indent = Pt(36)
            paragraph.paragraph_format.right_indent = Pt(36)


def style_exists(doc, style_name: str) -> bool:
    """
    Check if a style exists in the document.

    Parameters
    ----------
    doc : Document
        The document to check.
    style_name : str
        The name of the style.

    Returns
    -------
    bool
        True if the style exists, False otherwise.
    """
    return style_name in [style.name for style in doc.styles]


def set_outline_level(paragraph, level: int) -> None:
    """
    Set the outline level for a paragraph.

    Parameters
    ----------
    paragraph : Paragraph
        The paragraph to set the outline level for.
    level : int
        The outline level (0-9).
    """
    try:
        import docx

        # Get the paragraph properties element
        p_pr = paragraph._element.get_or_add_pPr()

        # Create the outline level element
        outline_lvl = docx.oxml.OxmlElement('w:outlineLvl')
        outline_lvl.set(docx.oxml.ns.qn('w:val'), str(level))

        # Remove any existing outline level
        for old_outline in p_pr.findall(docx.oxml.ns.qn('w:outlineLvl')):
            p_pr.remove(old_outline)

        # Add the new outline level
        p_pr.append(outline_lvl)
    except Exception:
        # Ignore errors in setting outline level
        pass
