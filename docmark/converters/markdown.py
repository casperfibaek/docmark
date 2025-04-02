"""
Markdown to DOCX converter for DocMark.

This module provides functionality for converting Markdown files to DOCX.
"""

import os
import re
import tempfile
from typing import Optional

import markdown
import docx
from docx import Document
from docx.shared import Pt, Inches
from bs4 import BeautifulSoup
from PIL import Image

from docmark.core.formatter import DocxFormatter


class MarkdownToDocxConverter:
    """Converter for Markdown to DOCX."""

    def __init__(
        self,
        formatter: Optional[DocxFormatter] = None,
        verbose: bool = False,
    ):
        """
        Initialize the Markdown to DOCX converter.

        Parameters
        ----------
        formatter : DocxFormatter, optional
            Formatter for DOCX styling.
        verbose : bool, optional
            Whether to print verbose output.
        """
        self.formatter = formatter or DocxFormatter()
        self.verbose = verbose

    def convert(
        self,
        markdown_path: str,
        output_path: str,
        template_path: Optional[str] = None,
        include_toc: bool = False,
    ) -> str:
        """
        Convert a Markdown file to DOCX.

        Parameters
        ----------
        markdown_path : str
            Path to the Markdown file.
        output_path : str
            Path to the output DOCX file.
        template_path : str, optional
            Path to the template file.
        include_toc : bool, optional
            Whether to include a table of contents.

        Returns
        -------
        str
            Path to the output DOCX file.
        """
        if self.verbose:
            print(f"Converting Markdown to DOCX: {markdown_path}")

        # Read the markdown content
        with open(markdown_path, "r", encoding="utf-8") as f:
            markdown_content = f.read()

        # Check for TOC marker
        has_toc_placeholder = '[TOC]' in markdown_content or '[[TOC]]' in markdown_content
        toc_position = None

        # Find the position of the TOC marker in the file
        if has_toc_placeholder:
            toc_marker_position = markdown_content.find('[TOC]')
            if toc_marker_position == -1:
                toc_marker_position = markdown_content.find('[[TOC]]')

            # Count the number of paragraphs and headers before TOC marker
            content_before_toc = markdown_content[:toc_marker_position]
            toc_position = len(re.findall(r'(?m)^#{1,6}\s|^[^\n]+\n', content_before_toc))

            if self.verbose:
                print(f"Found TOC marker at position {toc_position}")

            # Remove the TOC markers from the content
            markdown_content = markdown_content.replace('[TOC]', '')
            markdown_content = markdown_content.replace('[[TOC]]', '')

        # Create a new document with template styles if provided
        if template_path:
            if self.verbose:
                print(f"Using styles from template: {template_path}")
            doc = self._create_document_from_template(template_path)
        else:
            doc = Document()

        # Convert markdown to HTML
        html_content = markdown.markdown(
            markdown_content,
            extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.nl2br'
            ]
        )

        # Parse HTML using BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Process the HTML elements and add them to the document
        self._process_html_elements(
            soup,
            doc,
            os.path.dirname(markdown_path),
            toc_position=toc_position or (0 if include_toc else None)
        )

        # Save the document
        doc.save(output_path)

        if self.verbose:
            print(f"Saved DOCX to {output_path}")
            if has_toc_placeholder or include_toc:
                print("\nIMPORTANT: The Table of Contents will appear empty until updated.")
                print("To make the TOC appear in Word:")
                print("1. Open the document in Microsoft Word")
                print("2. Right-click the Table of Contents area")
                print("3. Select 'Update Field' or press F9")
                print("4. Choose 'Update entire table'")

        return output_path

    def _create_document_from_template(self, template_path: str) -> Document:
        """
        Create a new document based on a template.

        Parameters
        ----------
        template_path : str
            Path to the template file.

        Returns
        -------
        Document
            The created document.
        """
        try:
            # Create a document from the template directly
            result_doc = Document(template_path)

            # Clear all content, preserving only section properties
            for element in list(result_doc.element.body):
                if not element.tag.endswith('sectPr'):
                    result_doc.element.body.remove(element)

            if self.verbose:
                print("Using template with content safely removed")

            return result_doc

        except Exception as e:
            if self.verbose:
                print(f"Error using template: {str(e)}")
                print("Falling back to default document")

            # Fallback to default document
            return Document()

    def _process_html_elements(
        self,
        soup: BeautifulSoup,
        doc: Document,
        base_path: str,
        toc_position: Optional[int] = None,
    ) -> None:
        """
        Process HTML elements and add them to the Word document.

        Parameters
        ----------
        soup : BeautifulSoup
            BeautifulSoup object containing the HTML content.
        doc : Document
            The Word document to add content to.
        base_path : str
            Base path for resolving relative image paths.
        toc_position : int, optional
            Position (paragraph number) where TOC should be inserted.
        """
        # Process the body content
        elements = soup.body.children if soup.body else soup.children

        # Track the current position in the document
        current_position = 0

        for element in elements:
            if element.name is None:
                continue

            # Insert TOC at the specified position
            if toc_position is not None and current_position == toc_position:
                if self.verbose:
                    print(f"Inserting TOC at position {current_position}")
                self._add_table_of_contents(doc)

            current_position += 1

            if element.name == 'h1':
                heading = doc.add_heading(element.get_text(), level=1)
                self.formatter.apply_heading_style(heading, 1, doc)
            elif element.name == 'h2':
                heading = doc.add_heading(element.get_text(), level=2)
                self.formatter.apply_heading_style(heading, 2, doc)
            elif element.name == 'h3':
                heading = doc.add_heading(element.get_text(), level=3)
                self.formatter.apply_heading_style(heading, 3, doc)
            elif element.name == 'h4':
                heading = doc.add_heading(element.get_text(), level=4)
                self.formatter.apply_heading_style(heading, 4, doc)
            elif element.name == 'h5':
                heading = doc.add_heading(element.get_text(), level=5)
                self.formatter.apply_heading_style(heading, 5, doc)
            elif element.name == 'h6':
                heading = doc.add_heading(element.get_text(), level=6)
                self.formatter.apply_heading_style(heading, 6, doc)
            elif element.name == 'p':
                # Check if paragraph contains an image
                img = element.find('img')
                if img:
                    self._add_image_to_doc(doc, img, base_path)
                else:
                    # Handle text formatting (bold, italic, etc.)
                    para = doc.add_paragraph()
                    self._process_paragraph_text(para, element)
            elif element.name == 'ul':
                self._process_list(doc, element, is_numbered=False)
            elif element.name == 'ol':
                self._process_list(doc, element, is_numbered=True)
            elif element.name == 'table':
                self._process_table(doc, element)
            elif element.name == 'hr':
                doc.add_paragraph('').add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
            elif element.name == 'pre':
                # Code blocks
                code = element.get_text()
                p = doc.add_paragraph(code)
                self.formatter.apply_code_style(p, doc)

    def _process_paragraph_text(self, para, element):
        """Process paragraph text including formatting."""
        for child in element.children:
            if child.name is None:  # Plain text
                if child.string:  # Check if string is not None
                    run = para.add_run(child.string)
            elif child.name in ('strong', 'b'):
                run = para.add_run(child.get_text())
                run.bold = True
            elif child.name in ('em', 'i'):
                run = para.add_run(child.get_text())
                run.italic = True
            elif child.name == 'a':
                run = para.add_run(child.get_text())
                run.underline = True
                # Add hyperlink if possible
                if 'href' in child.attrs:
                    # Store hyperlink for reference
                    run._element.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = (
                        child['href']
                    )
            elif child.name == 'code':
                run = para.add_run(child.get_text())
                run.font.name = 'Courier New'
            elif child.name == 'br':
                para.add_run().add_break(docx.enum.text.WD_BREAK.LINE)
            else:
                # Handle nested formatting elements recursively
                self._process_paragraph_text(para, child)

    def _add_image_to_doc(self, doc, img_tag, base_path):
        """Add an image to the document."""
        if 'src' in img_tag.attrs:
            img_src = img_tag['src']

            # Handle relative paths
            if not os.path.isabs(img_src):
                if img_src.startswith('./'):
                    img_src = img_src[2:]
                img_path = os.path.join(base_path, img_src)
            else:
                img_path = img_src

            if os.path.exists(img_path):
                try:
                    # Get image dimensions with PIL
                    with Image.open(img_path) as pil_img:
                        # Check if image needs to be converted to RGB (for transparency issues)
                        if pil_img.mode in ('RGBA', 'LA') or (pil_img.mode == 'P' and 'transparency' in pil_img.info):
                            # Create a temp file for the converted image
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                                # Convert to RGB (white background) to handle transparency
                                bg = Image.new('RGB', pil_img.size, (255, 255, 255))
                                if pil_img.mode == 'RGBA':
                                    bg.paste(pil_img, mask=pil_img.split()[3])  # Use alpha channel as mask
                                else:
                                    bg.paste(pil_img)

                                # Save as PNG
                                bg.save(tmp.name)
                                img_path = tmp.name
                                if self.verbose:
                                    print(f"Converted image to RGB format: {img_path}")

                        # Get dimensions for scaling
                        width, height = pil_img.size

                    # Standard dpi for screen (typical conversion factor)
                    dpi = 96.0

                    # Calculate width in inches from pixels
                    width_inches = width / dpi
                    height_inches = height / dpi

                    # Calculate scaled dimensions to fit nicely in document
                    max_width_inches = 6.0  # Maximum width in inches

                    if width_inches > max_width_inches:
                        # Scale down while preserving aspect ratio
                        scale_factor = max_width_inches / width_inches
                        width_inches = max_width_inches
                        height_inches = height_inches * scale_factor

                    # Add image to document with proper size
                    image = doc.add_picture(img_path, width=Inches(width_inches), height=Inches(height_inches))

                    # Apply image style
                    self.formatter.apply_image_style(image, doc, width_inches=width_inches)

                    # Add image caption/alt text as a paragraph if present
                    if 'alt' in img_tag.attrs and img_tag['alt']:
                        caption = doc.add_paragraph(f"Figure: {img_tag['alt']}")
                        if hasattr(doc.styles, 'add_style') and 'Caption' not in [s.name for s in doc.styles]:
                            caption_style = doc.styles.add_style('Caption', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                            caption_style.font.italic = True
                            caption_style.font.size = Pt(10)
                            caption.style = caption_style
                        elif 'Caption' in [s.name for s in doc.styles]:
                            caption.style = doc.styles['Caption']

                    if self.verbose:
                        print(f"Added image: {img_path} (Size: {width_inches:.2f}\" x {height_inches:.2f}\")")

                    # Clean up temporary file if one was created
                    if img_path != os.path.join(base_path, img_src) and img_path.startswith(tempfile.gettempdir()):
                        try:
                            os.unlink(img_path)
                        except Exception:
                            pass

                except Exception as e:
                    if self.verbose:
                        print(f"Error adding image {img_path}: {str(e)}")
                    doc.add_paragraph(f"[Image: {img_path} - Error: {str(e)}]")
            else:
                if self.verbose:
                    print(f"Image not found: {img_path}")
                doc.add_paragraph(f"[Image not found: {img_path}]")

    def _process_list(self, doc, list_elem, is_numbered=False, level=1):
        """Process ordered and unordered lists."""
        for li in list_elem.find_all('li', recursive=False):
            # Add list item with proper style
            p = doc.add_paragraph()
            self.formatter.apply_list_style(p, is_numbered, level, doc)

            # Process text content of list item
            self._process_paragraph_text(p, li)

            # Process nested lists
            nested_ul = li.find('ul')
            nested_ol = li.find('ol')
            if nested_ul:
                self._process_list(doc, nested_ul, is_numbered=False, level=level+1)
            if nested_ol:
                self._process_list(doc, nested_ol, is_numbered=True, level=level+1)

    def _process_table(self, doc, table_elem):
        """Process tables."""
        rows = table_elem.find_all('tr')
        if not rows:
            return

        # Get number of columns based on first row
        first_row = rows[0]
        num_cols = len(first_row.find_all(['th', 'td']))

        # Create table
        table = doc.add_table(rows=len(rows), cols=num_cols)
        self.formatter.apply_table_style(table, doc)

        # Process rows
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            # Process cells
            for j, cell in enumerate(cells):
                if j < num_cols:  # Avoid index errors if row has more cells than expected
                    text = cell.get_text().strip()
                    table.cell(i, j).text = text

                    # Apply header formatting if this is a header cell
                    if cell.name == 'th' or i == 0:
                        for paragraph in table.cell(i, j).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _add_table_of_contents(self, doc):
        """
        Add a Table of Contents to the document.

        Parameters
        ----------
        doc : Document
            The Word document to add the TOC to.
        """
        try:
            # Add a title for the TOC
            toc_heading = doc.add_heading("Table of Contents", level=1)
            if 'TOC Heading' in [s.name for s in doc.styles]:
                toc_heading.style = doc.styles['TOC Heading']

            # Add a simple "update me" instruction
            instruction_para = doc.add_paragraph()
            run = instruction_para.add_run("(Press F9 to update the Table of Contents)")
            run.italic = True
            run.font.size = Pt(8)  # Smaller text

            # Add the TOC field code
            paragraph = doc.add_paragraph()
            paragraph._p.addnext(self._create_toc_field())

            # Add spacing after TOC
            doc.add_paragraph()

            if self.verbose:
                print("Added Table of Contents field")
                print("The document must be opened in Word and the TOC updated manually")

        except Exception as e:
            if self.verbose:
                print(f"Error adding Table of Contents: {str(e)}")

    def _create_toc_field(self):
        """Create a TOC field element using direct XML manipulation."""
        # Create the basic paragraph element that will contain our TOC
        p = docx.oxml.OxmlElement('w:p')

        # Create a run element
        r = docx.oxml.OxmlElement('w:r')
        p.append(r)

        # Create the TOC field using the OOXML field syntax
        field_char1 = docx.oxml.OxmlElement('w:fldChar')
        field_char1.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
        r.append(field_char1)

        # Field instruction - TOC with standard options
        instr_text = docx.oxml.OxmlElement('w:instrText')
        instr_text.set(docx.oxml.ns.qn('xml:space'), 'preserve')
        instr_text.text = ' TOC \\o "1-3" \\h \\z '
        r.append(instr_text)

        # Separating character
        field_char2 = docx.oxml.OxmlElement('w:fldChar')
        field_char2.set(docx.oxml.ns.qn('w:fldCharType'), 'separate')
        r.append(field_char2)

        # A minimal placeholder
        t = docx.oxml.OxmlElement('w:t')
        t.text = " "  # Just a space
        r.append(t)

        # End field character
        field_char3 = docx.oxml.OxmlElement('w:fldChar')
        field_char3.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
        r.append(field_char3)

        return p
